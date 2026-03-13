import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

INPUT = Path(__file__).resolve().parent.parent / "CURSOR-SETUP.md"
OUTPUT = Path.home() / "Downloads" / "Cursor-IDE-Full-Machine-Setup-Guide.docx"

md = INPUT.read_text()

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

for level in range(1, 5):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Calibri"
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

code_char_style = doc.styles.add_style("CodeChar", 2)  # character style
code_char_style.font.name = "Consolas"
code_char_style.font.size = Pt(9.5)
code_char_style.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)


def add_code_block(text):
    for i, line in enumerate(text.split("\n")):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)


def add_table(header_line, rows):
    cols = [c.strip() for c in header_line.strip().strip("|").split("|")]
    num_cols = len(cols)
    table = doc.add_table(rows=1, cols=num_cols, style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, col in enumerate(cols):
        cell = table.rows[0].cells[i]
        cell.text = re.sub(r"\*\*(.+?)\*\*", r"\1", col)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for row_text in rows:
        cells_text = [c.strip() for c in row_text.strip().strip("|").split("|")]
        row = table.add_row()
        for i in range(min(len(cells_text), num_cols)):
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", cells_text[i])
            clean = re.sub(r"`(.+?)`", r"\1", clean)
            clean = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", clean)
            row.cells[i].text = clean
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()


def add_rich_paragraph(text, heading_level=None, bold=False):
    if heading_level:
        p = doc.add_heading(level=heading_level)
    else:
        p = doc.add_paragraph()

    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*|\[([^\]]+)\]\([^)]+\))", text)
    for part in parts:
        if part is None:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
        elif part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("["):
            m = re.match(r"\[(.+?)\]\(.+?\)", part)
            if m:
                run = p.add_run(m.group(1))
                run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            else:
                p.add_run(part)
        else:
            run = p.add_run(part)
            if bold:
                run.bold = True
    return p


lines = md.split("\n")
i = 0
in_code = False
code_buffer = []

while i < len(lines):
    line = lines[i]

    # Code block start/end
    if line.startswith("```") or line.startswith("````"):
        if in_code:
            add_code_block("\n".join(code_buffer))
            code_buffer = []
            in_code = False
            i += 1
            continue
        else:
            in_code = True
            i += 1
            continue

    if in_code:
        code_buffer.append(line)
        i += 1
        continue

    # Skip HTML-ish lines
    if line.startswith(">"):
        text = line.lstrip("> ").strip()
        if text:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        i += 1
        continue

    # Horizontal rule
    if line.strip() == "---":
        doc.add_paragraph("─" * 70)
        i += 1
        continue

    # Headings
    heading_match = re.match(r"^(#{1,4})\s+(.+)", line)
    if heading_match:
        level = len(heading_match.group(1))
        text = heading_match.group(2)
        add_rich_paragraph(text, heading_level=level)
        i += 1
        continue

    # Table detection
    if "|" in line and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|", lines[i + 1]):
        header = line
        i += 2  # skip separator
        table_rows = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            table_rows.append(lines[i])
            i += 1
        add_table(header, table_rows)
        continue

    # Numbered list
    numbered = re.match(r"^(\d+)\.\s+(.+)", line)
    if numbered:
        add_rich_paragraph(f"{numbered.group(1)}. {numbered.group(2)}")
        i += 1
        continue

    # Bullet list
    if line.startswith("- "):
        text = line[2:].strip()
        add_rich_paragraph(f"  •  {text}")
        i += 1
        continue

    # Empty line
    if not line.strip():
        i += 1
        continue

    # Regular paragraph
    add_rich_paragraph(line)
    i += 1

doc.save(str(OUTPUT))
print(f"Saved to {OUTPUT}")
